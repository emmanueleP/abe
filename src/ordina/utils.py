from datetime import datetime
import os
from .settings import ordina_settings as settings
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_protocol():
    """Genera un numero di protocollo e timestamp."""
    timestamp = datetime.now()
    protocol_number = settings.get_next_protocol_number()
    timestamp_formatted = timestamp.strftime("%d/%m/%Y %H:%M:%S")
    return protocol_number, timestamp_formatted

def get_output_path(file_path):
    """Genera il percorso di output per il file protocollato."""
    filename = f"protocollato_{os.path.basename(file_path)}"
    output_dir = settings.get_output_directory()
    return os.path.join(output_dir, filename) 

def add_timestamp(doc):
    """Aggiunge il timbro al documento"""
    timestamp_para = doc.add_paragraph()
    timestamp_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    protocol_number = settings.get_next_protocol_number()
    timestamp_run = timestamp_para.add_run(
        f"Prot. N° {protocol_number}/{datetime.now().year}"
    )
    timestamp_run.font.name = "Arial"
    timestamp_run.font.size = Pt(13)
    timestamp_run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()
    return doc 