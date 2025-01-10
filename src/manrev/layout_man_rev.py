from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .images_manager import images_manager
from .settings import manrev_settings
import os

class DocumentLayout:
    def __init__(self, document):
        self.document = document
        self.sections = document.sections
        self.section = self.sections[0]
    
    def set_margins(self):
        """Imposta i margini del documento"""
        self.section.top_margin = Cm(2)
        self.section.bottom_margin = Cm(2)
        self.section.left_margin = Cm(2)
        self.section.right_margin = Cm(2)

    def add_header(self, title, year, number):
        """Aggiunge l'intestazione del documento"""
        # Aggiungi immagine sede se presente
        sede_path = manrev_settings.current_settings.get("sede_image", "")
        if sede_path and os.path.exists(sede_path):
            sede_para = self.document.add_paragraph()
            sede_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                sede_para.add_run().add_picture(sede_path, width=Cm(6))
                sede_para.add_run("\n")
            except Exception as e:
                print(f"Errore nel caricamento dell'immagine della sede: {e}")
        
        # Titolo e numero
        header = self.document.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = header.add_run(f"{title} N. {number}/{year}")
        title_run.bold = True
        title_run.font.size = Pt(14)

    def add_details_table(self, details):
        """Aggiunge la tabella dei dettagli"""
        table = self.document.add_table(rows=len(details), cols=2)
        table.style = 'Table Grid'
        
        for i, (key, value) in enumerate(details.items()):
            # Cella sinistra (chiave)
            cell = table.cell(i, 0)
            if key == 'Importo':
                cell.text = 'Importo in €'
            else:
                cell.text = key
            cell.paragraphs[0].runs[0].bold = True
            
            # Cella destra (valore)
            cell = table.cell(i, 1)
            if key == 'Importo':
                cell.text = f"{value} €"
            else:
                cell.text = str(value)

    def add_amount_text(self, amount_text):
        """Aggiunge l'importo in lettere"""
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = para.add_run(f"Importo in lettere: {amount_text}")
        run.bold = True
        run.font.size = Pt(11)

    def add_signatures(self, signatures):
        """Aggiunge le firme con layout migliorato"""
        # Aggiungi spazio prima delle firme
        self.document.add_paragraph()
        
        # Crea una tabella invisibile per allineare le firme
        table = self.document.add_table(rows=len(signatures), cols=1)
        table.allow_autofit = True
        
        # Rimuovi i bordi della tabella
        table.style = 'Table Grid'
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run._element.get_or_add_rPr().xpath('./w:bdr')
                        
        # Aggiungi le firme
        for i, (title, name) in enumerate(signatures.items()):
            cell = table.cell(i, 0)
            
            # Titolo
            title_para = cell.paragraphs[0]
            title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_run = title_para.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(11)
            
            # Nome
            name_para = cell.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            name_para.add_run(name).font.size = Pt(11)
            
            # Spazio per la firma
            signature_para = cell.add_paragraph()
            signature_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Aggiungi immagine firma se presente
            signature_path = images_manager.get_signature_path(title.lower())
            if signature_path:
                signature_para.add_run().add_picture(signature_path, width=Cm(4))
            else:
                # Se non c'è firma, aggiungi spazio per la firma manuale
                signature_para.add_run("_" * 40)
            
            # Aggiungi spazio dopo ogni firma
            cell.add_paragraph().add_run().add_break()
        
        # Aggiungi spazio dopo la tabella
        self.document.add_paragraph()

    def add_footer(self, place, date):
        """Aggiunge il piè di pagina"""
        footer = self.document.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
        footer.add_run(f"\n{place}, {date}") 