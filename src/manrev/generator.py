from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from win32com import client
from win32 import win32print
from PyQt5.QtWidgets import QMessageBox
from .settings import manrev_settings

def numero_in_parole(numero):
    """Converte un numero in parole in italiano"""
    unita = ["", "uno", "due", "tre", "quattro", "cinque", "sei", "sette", "otto", "nove"]
    decine = ["", "dieci", "venti", "trenta", "quaranta", "cinquanta", "sessanta", "settanta", "ottanta", "novanta"]
    teens = ["dieci", "undici", "dodici", "tredici", "quattordici", "quindici", "sedici", "diciassette", "diciotto", "diciannove"]
    
    try:
        euro = int(float(numero))
        centesimi = int((float(numero) * 100) % 100)
        
        risultato = []
        
        # Gestione euro
        if euro == 1:
            risultato.append("un euro")
        elif euro == 0:
            risultato.append("zero euro")
        else:
            if euro < 10:
                risultato.append(f"{unita[euro]} euro")
            elif euro < 20:
                risultato.append(f"{teens[euro-10]} euro")
            else:
                dec = euro // 10
                uni = euro % 10
                if uni == 0:
                    risultato.append(f"{decine[dec]} euro")
                else:
                    if uni == 1 and dec >= 2:
                        risultato.append(f"{decine[dec][:-1]}uno euro")
                    else:
                        risultato.append(f"{decine[dec]}{unita[uni]} euro")
        
        # Gestione centesimi
        if centesimi > 0:
            risultato.append("e")
            if centesimi < 10:
                risultato.append(f"{unita[centesimi]} centesimi")
            elif centesimi < 20:
                risultato.append(f"{teens[centesimi-10]} centesimi")
            else:
                dec = centesimi // 10
                uni = centesimi % 10
                if uni == 0:
                    risultato.append(f"{decine[dec]} centesimi")
                else:
                    if uni == 1 and dec >= 2:
                        risultato.append(f"{decine[dec][:-1]}uno centesimi")
                    else:
                        risultato.append(f"{decine[dec]}{unita[uni]} centesimi")
        else:
            risultato.extend(["e", "zero centesimi"])
        
        return " ".join(risultato)
        
    except Exception as e:
        raise Exception(f"Errore nella conversione del numero in parole: {str(e)}")

def generate_documents(data, file_path, print_after=False):
    """Genera il documento Word"""
    try:
        doc = Document()
        
        # Intestazione
        if manrev_settings.current_settings["header_image"]:
            try:
                doc.add_picture(
                    manrev_settings.current_settings["header_image"], 
                    width=Inches(2)
                )
            except Exception as e:
                print(f"Errore nel caricamento dell'immagine di intestazione: {e}")
        
        # Titolo
        heading = doc.add_heading(f"AVIS Comunale Decimoputzu OdV - {data['Tipo']}", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Numero documento
        p = doc.add_paragraph()
        p.add_run(f"Tipo documento: {data['Tipo']}").bold = True
        p.add_run(f" N° {data['Numero']}").bold = True
        
        # Capitolo
        p = doc.add_paragraph()
        p.add_run("Capitolo: ").bold = True
        p.add_run(data['Capitolo']).bold = True
        
        # Importo
        importo_str = str(data['Importo in €'])
        importo_numerico = float(importo_str.replace(',', '.'))
        importo_display = str(importo_numerico).replace('.', ',')
        importo_in_lettere = numero_in_parole(importo_numerico)
        
        p = doc.add_paragraph()
        p.add_run("Importo in €: ").bold = True
        p.add_run(importo_display).bold = True
        doc.add_paragraph(f"Importo in lettere: {importo_in_lettere}")
        
        # Descrizione
        doc.add_paragraph("Descrizione del pagamento:").bold = True
        p = doc.add_paragraph(data['Descrizione del pagamento'])
        p_format = p.paragraph_format
        p_format.left_indent = Inches(0.5)
        p_format.right_indent = Inches(0.5)
        p_format.space_before = Pt(12)
        p_format.space_after = Pt(12)
        
        # Luogo e data
        doc.add_paragraph(f"Luogo: {data['Luogo']}")
        doc.add_paragraph(f"Data: {data['Data']}")
        
        # Firme
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
        
        for title, name, img_path in [
            ("Il Tesoriere", data['Il Tesoriere'], manrev_settings.current_settings["firme"]["tesoriere_firma"]),
            ("Il Presidente", data['Il Presidente'], manrev_settings.current_settings["firme"]["presidente_firma"]),
            ("L'Addetto Contabile", data["L'Addetto Contabile"], manrev_settings.current_settings["firme"]["addetto_firma"])
        ]:
            run = p.add_run(title + "\n")
            run.bold = True
            
            if img_path and os.path.exists(img_path):
                p.add_run().add_picture(img_path, width=Inches(1))
            
            p.add_run("\n" + name + "\t\t")
        
        # Salva il documento
        doc.save(file_path)
        
        # Stampa se richiesto
        if print_after:
            print_document(file_path)
            
        # Apri il documento
        os.startfile(file_path)
        
        return file_path
        
    except Exception as e:
        raise Exception(f"Errore durante la generazione del documento: {str(e)}")

def print_document(file_path):
    """Stampa il documento"""
    try:
        shell = client.Dispatch("Shell.Application")
        shell.ShellExecute(0, "print", file_path, None, ".", 0)
        return True
    except Exception as e:
        raise Exception(f"Errore durante la stampa: {str(e)}")