import os
from PIL import Image, ImageDraw, ImageFont
from PyPDF2 import PdfWriter, PdfReader
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook
from openpyxl.drawing.text import Font
from openpyxl.styles import PatternFill, Font as XLFont, Alignment
from openpyxl import drawing
from openpyxl.utils import get_column_letter
import io
from utils import generate_protocol, get_output_path
from settings import settings
import json
from datetime import datetime

def handle_file(file_path):
    """
    Gestisce la protocollazione del file in base al suo formato.
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError("Il file non esiste")

        protocol_number, timestamp = generate_protocol()
        
        if file_path.lower().endswith((".png", ".jpg", ".jpeg")):
            result = add_stamp_to_image(file_path, protocol_number, timestamp)
        elif file_path.lower().endswith(".pdf"):
            result = add_stamp_to_pdf(file_path, protocol_number, timestamp)
        elif file_path.lower().endswith(".docx"):
            result = add_stamp_to_docx(file_path, protocol_number, timestamp)
        elif file_path.lower().endswith(".xlsx"):
            result = add_stamp_to_xlsx(file_path, protocol_number, timestamp)
        else:
            return "Formato file non supportato."
        
        # Salva nella cronologia
        save_to_history(protocol_number, get_output_path(file_path))
        
        return result
            
    except Exception as e:
        raise Exception(f"Errore durante la protocollazione: {str(e)}")

def add_stamp_to_docx(file_path, protocol, timestamp):
    """
    Aggiunge il timbro di protocollo al documento Word.
    """
    try:
        # Carica il documento
        doc = Document(file_path)
        
        # Aggiungi una sezione all'inizio del documento
        section = doc.sections[0]
        header = section.header
        
        # Aggiungi il testo del protocollo nell'header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = paragraph.add_run(f"{protocol}\n{timestamp}")
        font = run.font
        font.name = 'Arial'
        font.size = Pt(12)
        font.color.rgb = RGBColor(255, 0, 0)
        
        # Aggiungi il timbro personalizzato se presente
        stamp_img = settings.get_stamp_image()
        if stamp_img:
            # Salva temporaneamente il timbro
            temp_stamp = "temp_stamp.png"
            stamp_img.save(temp_stamp)
            
            # Aggiungi il timbro al piè di pagina
            footer = section.footer
            footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = footer_paragraph.add_run()
            run.add_picture(temp_stamp, width=Inches(1))
            
            # Rimuovi il file temporaneo
            os.remove(temp_stamp)
        
        # Salva il documento
        output_path = get_output_path(file_path)
        doc.save(output_path)
        
        return f"Documento protocollato con successo!\nNumero protocollo: {protocol}\nSalvato in: {output_path}"
        
    except Exception as e:
        raise Exception(f"Errore durante la protocollazione del documento Word: {str(e)}")

def add_stamp_to_xlsx(file_path, protocol, timestamp):
    """
    Aggiunge il timbro di protocollo al foglio Excel.
    """
    try:
        # Carica il workbook
        wb = load_workbook(file_path)
        ws = wb.active
        
        # Aggiungi il testo del protocollo in alto a sinistra
        stamp_text = f"{protocol}\n{timestamp}"
        ws['A1'] = stamp_text
        
        # Formatta la cella
        cell = ws['A1']
        cell.font = XLFont(name='Arial', size=12, color='FF0000')
        cell.alignment = Alignment(wrap_text=True)
        ws.row_dimensions[1].height = 40
        
        # Aggiungi il timbro personalizzato se presente
        stamp_img = settings.get_stamp_image()
        if stamp_img:
            # Salva temporaneamente il timbro
            temp_stamp = "temp_stamp.png"
            stamp_img.save(temp_stamp)
            
            # Calcola l'ultima riga e colonna usate
            max_row = ws.max_row
            max_col = ws.max_column
            
            # Aggiungi l'immagine in basso a destra
            img = drawing.image.Image(temp_stamp)
            img.width = 100
            img.height = 100
            img.anchor = f"{get_column_letter(max_col)}_{max_row}"
            ws.add_image(img)
            
            # Rimuovi il file temporaneo
            os.remove(temp_stamp)
        
        # Salva il file
        output_path = get_output_path(file_path)
        wb.save(output_path)
        
        return f"Documento protocollato con successo!\nNumero protocollo: {protocol}\nSalvato in: {output_path}"
        
    except Exception as e:
        raise Exception(f"Errore durante la protocollazione del foglio Excel: {str(e)}")

def add_stamp_to_image(file_path, protocol, timestamp):
    """
    Aggiunge il timbro di protocollo all'immagine.
    """
    try:
        with Image.open(file_path) as img:
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            draw = ImageDraw.Draw(img)
            
            # Carica il font
            try:
                font = ImageFont.truetype("arial.ttf", 36)
            except:
                font = ImageFont.load_default()
            
            # Prepara il testo del timbro
            stamp_text = f"{protocol}\n{timestamp}"
            
            # Aggiungi il testo del protocollo
            margin = 20
            position = (margin, margin)
            
            # Aggiungi sfondo semi-trasparente
            text_bbox = draw.textbbox(position, stamp_text, font=font)
            padding = 10
            draw.rectangle(
                [
                    text_bbox[0] - padding,
                    text_bbox[1] - padding,
                    text_bbox[2] + padding,
                    text_bbox[3] + padding
                ],
                fill=(255, 255, 255, 128)
            )
            
            draw.text(position, stamp_text, fill=(255, 0, 0), font=font)
            
            # Aggiungi il timbro personalizzato
            stamp_img = settings.get_stamp_image()
            if stamp_img:
                # Posiziona il timbro in basso a destra
                position = (
                    img.width - stamp_img.width - margin,
                    img.height - stamp_img.height - margin
                )
                img.paste(stamp_img, position, stamp_img)
            
            output_path = get_output_path(file_path)
            img.save(output_path, quality=95)
            
            return f"Documento protocollato con successo!\nNumero protocollo: {protocol}\nSalvato in: {output_path}"
            
    except Exception as e:
        raise Exception(f"Errore durante la protocollazione dell'immagine: {str(e)}")

def add_stamp_to_pdf(file_path, protocol, timestamp):
    """
    Aggiunge il timbro di protocollo al PDF.
    """
    try:
        # Crea un'immagine con il timbro
        stamp_img = create_stamp_image(protocol, timestamp)
        
        # Apri il PDF
        reader = PdfReader(file_path)
        writer = PdfWriter()
        
        # Processa ogni pagina
        for i, page in enumerate(reader.pages):
            # Aggiungi il timbro solo alla prima pagina
            if i == 0:
                page.merge_page(stamp_img)
            writer.add_page(page)
        
        # Salva il PDF
        output_path = get_output_path(file_path)
        with open(output_path, "wb") as output_file:
            writer.write(output_file)
            
        return f"Documento protocollato con successo!\nNumero protocollo: {protocol}\nSalvato in: {output_path}"
        
    except Exception as e:
        raise Exception(f"Errore durante la protocollazione del PDF: {str(e)}")

def create_stamp_image(protocol, timestamp):
    """
    Crea un'immagine del timbro per il PDF.
    """
    # Dimensioni e margini
    width = 600
    height = 200
    margin = 20
    
    # Crea un'immagine con sfondo trasparente
    img = Image.new('RGBA', (width, height), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    
    try:
        font = ImageFont.truetype("arial.ttf", 24)
    except:
        font = ImageFont.load_default()
    
    # Aggiungi il testo del protocollo
    stamp_text = f"{protocol}\n{timestamp}"
    draw.text((margin, margin), stamp_text, fill=(255, 0, 0), font=font)
    
    # Aggiungi il timbro personalizzato se presente
    stamp_img = settings.get_stamp_image()
    if stamp_img:
        # Posiziona il timbro a destra
        position = (width - stamp_img.width - margin, margin)
        img.paste(stamp_img, position, stamp_img)
    
    # Converti l'immagine in PDF
    pdf_bytes = io.BytesIO()
    img.save(pdf_bytes, format='PDF')
    pdf_bytes.seek(0)
    
    return PdfReader(pdf_bytes).pages[0]

def save_to_history(protocol, file_path):
    """Salva l'operazione nella cronologia"""
    try:
        history_file = 'protocol_history.json'
        
        # Carica la cronologia esistente o crea una nuova
        if os.path.exists(history_file):
            try:
                with open(history_file, 'r', encoding='utf-8') as f:
                    history = json.load(f)
            except json.JSONDecodeError:
                history = {}
        else:
            history = {}
        
        # Ottieni l'anno corrente
        current_year = str(datetime.now().year)
        if current_year not in history:
            history[current_year] = []
        
        # Aggiungi la nuova voce
        entry = {
            'number': protocol,
            'date': datetime.now().strftime('%d/%m/%Y'),
            'time': datetime.now().strftime('%H:%M:%S'),
            'file': os.path.abspath(file_path)  # Usa il percorso assoluto
        }
        
        # Aggiungi in testa alla lista (i più recenti prima)
        history[current_year].insert(0, entry)
        
        # Salva la cronologia aggiornata
        with open(history_file, 'w', encoding='utf-8') as f:
            json.dump(history, f, indent=4, ensure_ascii=False)
            
    except Exception as e:
        print(f"Errore nel salvare la cronologia: {e}")
